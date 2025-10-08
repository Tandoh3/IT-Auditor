import streamlit as st
import pandas as pd
import zipfile
import rarfile
import io
import re
from collections import defaultdict
import matplotlib.pyplot as plt
import seaborn as sns
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import tempfile
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import json
from datetime import timedelta
from thefuzz import process, fuzz

# 🎨 Configure Streamlit Page
# --- Page Configuration ---
st.set_page_config(
    page_title="🔐 Your-IT-Auditor",
    page_icon="💻",
    layout="centered",  # 👈 now it's narrow
    initial_sidebar_state="expanded"
)

# --- Custom Styling for Width Control ---
st.markdown(
    """
    <style>
    .block-container {
        max-width: 1000px;  
        padding-top: 2rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# 🌟 Sidebar Navigation
st.sidebar.title("🔍 **Your-IT-Auditor**")
st.sidebar.markdown("---")

# Main category selection
main_category = st.sidebar.radio(
    "📌 Select Category:",
    ["🔐 Network Security", "👥 Identity & Access Management"]
)

st.sidebar.markdown("---")

# Sub-pages based on main category
if main_category == "🔐 Network Security":
    page = st.sidebar.radio(
        "🔧 Network Security Tools:",
        ["📅 Audit Planner", "📊 Config Audit"]
    )
else:  # Identity & Access Management
    page = st.sidebar.radio(
        "👤 IAM Tools:",
        ["🏠 IAM Main", "🔁 Duplicate User Provisioning", "📂 Database Groups", "🔑 Database Privilege Users", "🗂 Database Profiles"]
    )

st.sidebar.markdown("---")
st.sidebar.info("**Use this comprehensive tool to manage network security and user access!**")

# =============================================================================
# NETWORK SECURITY FUNCTIONS
# =============================================================================

# ---------------------------
# Audit Planning Assistant
# ---------------------------
def audit_planner():
    st.header("📅 Network Audit Planning Assistant")
    
    # Audit Scope
    st.subheader("1. Audit Scope Definition")
    col1, col2 = st.columns(2)
    
    with col1:
        audit_name = st.text_input("Audit Name", "Q1 2024 Network Security Audit")
        audit_type = st.selectbox(
            "Audit Type",
            ["Comprehensive Security", "Compliance Check", "Pre-Migration", "Post-Change", "Routine Maintenance"]
        )
    
    with col2:
        priority = st.select_slider("Priority Level", ["Low", "Medium", "High", "Critical"])
        timeline_days = st.number_input("Timeline (days)", min_value=1, max_value=90, value=14)
    
    # Device Inventory
    st.subheader("2. Device Inventory")
    
    device_types = st.multiselect(
        "Device Types to Audit",
        ["Routers", "Switches", "Firewalls", "Wireless Controllers", "Load Balancers", "VPN Gateways"],
        default=["Routers", "Switches", "Firewalls"]
    )
    
    estimated_devices = st.number_input("Estimated Number of Devices", min_value=1, max_value=1000, value=50)
    
    # Risk Assessment
    st.subheader("3. Risk Assessment Factors")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        business_impact = st.selectbox(
            "Business Impact",
            ["Low", "Medium", "High", "Critical"],
            help="Impact on business operations if devices fail"
        )
    
    with col2:
        data_sensitivity = st.selectbox(
            "Data Sensitivity",
            ["Public", "Internal", "Confidential", "Restricted"],
            help="Sensitivity of data handled by these devices"
        )
    
    with col3:
        compliance_requirements = st.multiselect(
            "Compliance Requirements",
            ["PCI-DSS", "HIPAA", "SOX", "GDPR", "NIST", "ISO 27001", "None"]
        )
    
    # Resource Planning
    st.subheader("4. Resource Planning")
    
    col1, col2 = st.columns(2)
    
    with col1:
        team_size = st.number_input("Team Size", min_value=1, max_value=20, value=3)
        hours_per_device = st.slider("Estimated Hours per Device", 0.5, 8.0, 2.0)
    
    with col2:
        expertise_level = st.selectbox(
            "Required Expertise Level",
            ["Junior", "Mid-Level", "Senior", "Expert"]
        )
        tools_available = st.multiselect(
            "Available Tools",
            ["Network Scanner", "Config Manager", "SIEM", "Vulnerability Scanner", "Custom Scripts"]
        )
    
    # Timeline Planning
    st.subheader("5. Timeline & Milestones")
    
    start_date = st.date_input("Planned Start Date", datetime.now() + timedelta(days=7))
    
    # Calculate timeline
    total_hours = estimated_devices * hours_per_device
    total_days = max(1, total_hours / (team_size * 8))  # 8 hours per day per person
    
    # Key milestones - convert dates to strings for JSON serialization
    milestones = {
        "Planning & Scoping": start_date.strftime("%Y-%m-%d"),
        "Data Collection": (start_date + timedelta(days=2)).strftime("%Y-%m-%d"),
        "Configuration Analysis": (start_date + timedelta(days=int(total_days * 0.3))).strftime("%Y-%m-%d"),
        "Vulnerability Assessment": (start_date + timedelta(days=int(total_days * 0.6))).strftime("%Y-%m-%d"),
        "Reporting": (start_date + timedelta(days=int(total_days * 0.8))).strftime("%Y-%m-%d"),
        "Remediation Planning": (start_date + timedelta(days=int(total_days))).strftime("%Y-%m-%d")
    }
    
    # Display planning summary
    if st.button("Generate Audit Plan"):
        st.success("🎯 Audit Plan Generated Successfully!")
        
        # Summary Section
        st.subheader("📋 Audit Plan Summary")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Devices", estimated_devices)
            st.metric("Team Size", team_size)
            st.metric("Priority", priority)
        
        with col2:
            st.metric("Total Hours", f"{total_hours:.1f}")
            st.metric("Estimated Days", f"{total_days:.1f}")
            st.metric("Business Impact", business_impact)
        
        with col3:
            st.metric("Start Date", start_date.strftime("%Y-%m-%d"))
            st.metric("End Date", (start_date + timedelta(days=total_days)).strftime("%Y-%m-%d"))
            st.metric("Risk Level", "High" if business_impact in ["High", "Critical"] else "Medium")
        
        # Timeline Visualization
        st.subheader("⏰ Project Timeline")
        timeline_data = []
        for milestone, date_str in milestones.items():
            date = datetime.strptime(date_str, "%Y-%m-%d").date()
            timeline_data.append({
                "Milestone": milestone,
                "Date": date_str,
                "Days from Start": (date - start_date).days
            })
        
        timeline_df = pd.DataFrame(timeline_data)
        st.dataframe(timeline_df, width='stretch')
        
        # Resource Allocation
        st.subheader("👥 Resource Allocation")
        
        resource_data = {
            "Task": ["Planning", "Data Collection", "Analysis", "Reporting", "Remediation"],
            "Effort (%)": [10, 25, 40, 15, 10],
            "Team Members": [team_size, team_size, team_size, team_size - 1, team_size]
        }
        resource_df = pd.DataFrame(resource_data)
        st.dataframe(resource_df, width='stretch')
        
        # Risk Matrix
        st.subheader("🚨 Risk Assessment Matrix")
        
        risk_matrix = {
            "Factor": ["Business Impact", "Data Sensitivity", "Compliance", "Team Expertise", "Tool Availability"],
            "Level": [business_impact, data_sensitivity, 
                     "High" if compliance_requirements else "Low", 
                     expertise_level,
                     "High" if len(tools_available) >= 3 else "Medium"],
            "Mitigation": [
                "Ensure backup systems available",
                "Focus on encryption & access controls",
                "Document compliance evidence",
                "Provide training if needed",
                "Plan for manual processes"
            ]
        }
        risk_df = pd.DataFrame(risk_matrix)
        st.dataframe(risk_df, width='stretch')
        
        # Export Plan
        st.subheader("📤 Export Audit Plan")
        
        audit_plan = {
            "audit_name": audit_name,
            "audit_type": audit_type,
            "priority": priority,
            "timeline_days": timeline_days,
            "device_types": device_types,
            "estimated_devices": estimated_devices,
            "team_size": team_size,
            "total_hours": total_hours,
            "start_date": start_date.strftime("%Y-%m-%d"),
            "milestones": milestones,
            "risk_factors": risk_matrix
        }
        
        # Word export
        word_bytes = generate_audit_plan_word(audit_plan)
        st.download_button(
            label="📄 Download Audit Plan (Word)",
            data=word_bytes,
            file_name=f"audit_plan_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # JSON as optional backup
        json_plan = json.dumps(audit_plan, indent=2)
        st.download_button(
            label="📊 Download Audit Plan (JSON)",
            data=json_plan,
            file_name=f"audit_plan_{datetime.now().strftime('%Y%m%d')}.json",
            mime="application/json"
        )
        
        # Text summary export
        text_summary = f"""
AUDIT PLAN: {audit_name}
==================================
Type: {audit_type}
Priority: {priority}
Timeline: {timeline_days} days
Start Date: {start_date.strftime('%Y-%m-%d')}

DEVICE INVENTORY:
-----------------
Types: {', '.join(device_types)}
Estimated Devices: {estimated_devices}

RESOURCE PLANNING:
------------------
Team Size: {team_size}
Expertise Level: {expertise_level}
Total Effort: {total_hours} hours
Estimated Duration: {total_days:.1f} days

RISK ASSESSMENT:
----------------
Business Impact: {business_impact}
Data Sensitivity: {data_sensitivity}
Compliance: {', '.join(compliance_requirements) if compliance_requirements else 'None'}

MILESTONES:
-----------
{chr(10).join([f'{milestone}: {date_str}' for milestone, date_str in milestones.items()])}
        """
        
        st.download_button(
            label="📝 Download Audit Summary (TXT)",
            data=text_summary,
            file_name=f"audit_summary_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )

# ---------------------------
# Word Document generator for Audit Plan
# ---------------------------
def generate_audit_plan_word(audit_plan):
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Title
    title = doc.add_heading('Network Audit Plan', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    # 1. Audit Overview
    doc.add_heading('1. Audit Overview', level=1)
    overview_table = doc.add_table(rows=4, cols=2)
    overview_table.style = 'Table Grid'
    
    overview_data = [
        ["Audit Name", audit_plan['audit_name']],
        ["Audit Type", audit_plan['audit_type']],
        ["Priority Level", audit_plan['priority']],
        ["Timeline (days)", str(audit_plan['timeline_days'])]
    ]
    
    for i, (label, value) in enumerate(overview_data):
        cells = overview_table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # 2. Device Inventory
    doc.add_heading('2. Device Inventory', level=1)
    doc.add_paragraph(f"Device Types: {', '.join(audit_plan['device_types'])}")
    doc.add_paragraph(f"Estimated Number of Devices: {audit_plan['estimated_devices']}")
    doc.add_paragraph()
    
    # 3. Resource Planning
    doc.add_heading('3. Resource Planning', level=1)
    resource_table = doc.add_table(rows=4, cols=2)
    resource_table.style = 'Table Grid'
    
    resource_data = [
        ["Team Size", str(audit_plan['team_size'])],
        ["Total Hours", f"{audit_plan['total_hours']:.1f}"],
        ["Estimated Duration", f"{audit_plan['total_hours'] / (audit_plan['team_size'] * 8):.1f} days"],
        ["Expertise Level", audit_plan.get('expertise_level', 'Not specified')]
    ]
    
    for i, (label, value) in enumerate(resource_data):
        cells = resource_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
    
    doc.add_paragraph()
    
    # 4. Timeline & Milestones
    doc.add_heading('4. Timeline & Milestones', level=1)
    milestones_table = doc.add_table(rows=len(audit_plan['milestones']) + 1, cols=2)
    milestones_table.style = 'Table Grid'
    
    # Header
    header_cells = milestones_table.rows[0].cells
    header_cells[0].text = "Milestone"
    header_cells[1].text = "Target Date"
    
    # Data
    for i, (milestone, date) in enumerate(audit_plan['milestones'].items(), 1):
        cells = milestones_table.rows[i].cells
        cells[0].text = milestone
        cells[1].text = date
    
    doc.add_paragraph()
    
    # 5. Risk Assessment
    doc.add_heading('5. Risk Assessment', level=1)
    risk_table = doc.add_table(rows=len(audit_plan['risk_factors']['Factor']) + 1, cols=3)
    risk_table.style = 'Table Grid'
    
    # Header
    risk_header = risk_table.rows[0].cells
    risk_header[0].text = "Factor"
    risk_header[1].text = "Level"
    risk_header[2].text = "Mitigation"
    
    # Data
    for i in range(len(audit_plan['risk_factors']['Factor'])):
        cells = risk_table.rows[i + 1].cells
        cells[0].text = audit_plan['risk_factors']['Factor'][i]
        cells[1].text = audit_plan['risk_factors']['Level'][i]
        cells[2].text = audit_plan['risk_factors']['Mitigation'][i]
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    word_bytes = buffer.getvalue()
    buffer.close()
    
    return word_bytes

# ---------------------------
# Network Config Audit Functions (Existing Code)
# ---------------------------
def audit_config(filename, content):
    findings = []
    # ... (include all your existing audit_config function code here)
    # Normalize content for easier regex
    # (we'll still use case-insensitive but keep content as-is)
    # --- 1. Layer 2 Security ---
    if not re.search(r"\bip dhcp snooping\b", content, re.IGNORECASE):
        findings.append(("DHCP Snooping Disabled", filename, "DHCP attacks possible", "Enable DHCP Snooping", "Layer 2"))

    if not re.search(r"\bip arp inspection\b", content, re.IGNORECASE):
        findings.append(("Dynamic ARP Inspection Missing", filename, "ARP spoofing possible", "Enable Dynamic ARP Inspection", "Layer 2"))

    if not re.search(r"\bswitchport port-security\b", content, re.IGNORECASE):
        findings.append(("Port Security Not Configured", filename, "MAC flooding risk", "Enable Port Security", "Layer 2"))

    # Heuristic: detect interface blocks that may not include shutdown
    try:
        # find interface blocks; heuristic: interface <name> ... (if no 'shutdown' in block, flag)
        iface_blocks = re.findall(r'(?ms)^(interface\s+\S+.*?)(?=^interface\s+\S+|\Z)', content, re.IGNORECASE)
        for block in iface_blocks:
            if not re.search(r'(?m)^\s*shutdown\b', block):
                # don't spam for each interface; append once per file as heuristic
                findings.append(("Unused Interfaces Active (heuristic)", filename, "Potential unused interfaces not administratively shutdown", "Review & administratively shutdown unused interfaces", "Layer 2"))
                break
    except Exception:
        pass

    if re.search(r"\bswitchport trunk native vlan\s+1\b", content, re.IGNORECASE):
        findings.append(("Default Native VLAN in Use", filename, "VLAN hopping risk", "Change native VLAN from 1", "Layer 2"))

    # --- 2. Access Control ---
    # Telnet detection across vty or transport input
    if re.search(r'(?mi)^\s*transport input .*telnet', content) or re.search(r'(?ms)^line vty.*?transport input .*telnet', content):
        findings.append(("Telnet Enabled", filename, "Credentials exposed in cleartext", "Disable Telnet and use SSH only", "Access Control"))

    if re.search(r"\bsnmp-server community\s+(public|private)\b", content, re.IGNORECASE):
        findings.append(("Default SNMP Community", filename, "Unauthorized SNMP access risk", "Use SNMPv3 with strong credentials", "Access Control"))

    if not re.search(r"\b(access-list|ip access-list|ip prefix-list|ipv6 access-list)\b", content, re.IGNORECASE):
        findings.append(("No ACLs Found", filename, "Unrestricted traffic flows", "Implement ACLs where needed", "Access Control"))

    # --- 3. Authentication & Authorization ---
    if not re.search(r"\baaa new-model\b", content, re.IGNORECASE):
        findings.append(("No AAA Configured", filename, "No centralized authentication", "Enable AAA (TACACS+/RADIUS)", "AAA"))

    if re.search(r'(?mi)^\s*username\s+\S+\s+(?:password|privilege)\b', content):
        findings.append(("Local User Accounts with Passwords", filename, "Local credential management; possible weak auth", "Use AAA and avoid plaintext local passwords", "AAA"))

    # --- 4. Logging & Monitoring ---
    if not re.search(r"\blogging\s+\S+", content, re.IGNORECASE):
        findings.append(("No Syslog Configured", filename, "No centralized log collection", "Configure Syslog servers", "Logging"))

    if not re.search(r"\b(ntp server|clock set|ntp peer)\b", content, re.IGNORECASE):
        findings.append(("No NTP Configured", filename, "Logs not time-synced", "Configure NTP servers", "Logging"))

    if not re.search(r"snmp-server group .* v3", content, re.IGNORECASE):
        findings.append(("SNMPv3 Not Configured", filename, "Monitoring unencrypted", "Use SNMPv3 with authentication & privacy", "Logging"))

    # --- 5. Cryptographic & Protocol Risks ---
    if re.search(r'(?mi)^\s*(service ftp|ftp server|ip ftp)\b', content):
        findings.append(("FTP Enabled", filename, "Credentials exposed in cleartext", "Disable FTP; use SFTP/SCP/FTPS", "Crypto"))

    if re.search(r'(?mi)^\s*ip http\b', content):
        findings.append(("HTTP Server Enabled", filename, "Management traffic unencrypted", "Disable HTTP; enable HTTPS (ip http secure-server)", "Crypto"))

    if not re.search(r"\bip ssh\b", content, re.IGNORECASE):
        findings.append(("SSH Not Configured", filename, "Secure remote management not enforced", "Enable SSH v2 and restrict vty to SSH", "Crypto"))

    # --- 6. Resilience & Availability ---
    if not re.search(r"\b(standby\b|vrrp\b|hsrp\b)", content, re.IGNORECASE):
        findings.append(("No First-Hop Redundancy (HSRP/VRRP)", filename, "Single point of failure for gateway", "Implement HSRP/VRRP where required", "Resilience"))

    if not re.search(r"\bstorm-control\b", content, re.IGNORECASE):
        findings.append(("No Storm Control", filename, "Broadcast/multicast flood risk", "Enable storm-control on access ports", "Resilience"))

    if not re.search(r"\bspanning-tree\b", content, re.IGNORECASE):
        findings.append(("Spanning Tree Not Configured", filename, "Switching loops possible", "Enable STP and configure root guard/portfast", "Resilience"))

    # --- 7. Configuration Management ---
    if re.search(r"\bpassword 7\b", content, re.IGNORECASE):
        findings.append(("Weak Password Encryption (Type 7)", filename, "Easily reversible encryption", "Avoid type 7; use enable secret / stronger hashes", "Config Mgmt"))

    if not re.search(r"\barchive\b", content, re.IGNORECASE):
        findings.append(("No Config Archiving", filename, "No config backup/versioning", "Enable config archive/backup/versioning", "Config Mgmt"))

    if not re.search(r"\bservice password-encryption\b", content, re.IGNORECASE):
        findings.append(("Passwords Not Encrypted", filename, "Plaintext passwords in config", "Enable 'service password-encryption' and use secrets", "Config Mgmt"))

    return findings

def get_risk_score(num_findings):
    if num_findings == 0:
        return "No Risk"
    elif num_findings <= 2:
        return "Low"
    elif num_findings <= 5:
        return "Medium"
    else:
        return "High"

def generate_heatmap_figure(df_findings):
    """Return matplotlib figure of heatmap (devices x categories counts)."""
    if df_findings.empty:
        fig = plt.figure(figsize=(6, 3))
        plt.text(0.5, 0.5, "No data", ha='center', va='center')
        plt.axis('off')
        return fig

    pivot = pd.pivot_table(df_findings, values='Finding', index='File', columns='Category', aggfunc='count', fill_value=0)
    # ensure consistent category order
    categories_order = ["Layer 2", "Access Control", "AAA", "Logging", "Crypto", "Resilience", "Config Mgmt"]
    cols = [c for c in categories_order if c in pivot.columns] + [c for c in pivot.columns if c not in categories_order]
    pivot = pivot[cols]
    fig, ax = plt.subplots(figsize=(10, max(2, 0.35 * len(pivot.index))))
    sns.heatmap(pivot, cmap="RdYlGn_r", annot=True, fmt="d", linewidths=0.5, ax=ax)
    ax.set_title("Risk Heatmap per Category (device = row)")
    plt.tight_layout()
    return fig

def generate_pdf_report(summary_df, df_findings, risk_counts, category_counts):
    # ... (include all your existing PDF generation code here)
    buffer = io.BytesIO()
    
    # Use landscape orientation for wider tables
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), 
                          leftMargin=0.3*inch, rightMargin=0.3*inch, 
                          topMargin=0.4*inch, bottomMargin=0.4*inch)
    
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        spaceAfter=12,
        alignment=1
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8
    )
    
    device_style = ParagraphStyle(
        'DeviceStyle',
        parent=styles['Heading3'],
        fontSize=10,
        spaceAfter=6,
        textColor=colors.darkblue
    )
    
    table_style = ParagraphStyle(
        'TableStyle',
        parent=styles['Normal'],
        fontSize=7,
        leading=8,
        spaceAfter=0,
        spaceBefore=0
    )
    
    elements = []

    # Title
    elements.append(Paragraph("Network Configuration Audit Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%SZ')}", table_style))
    elements.append(Spacer(1, 15))

    # 1. Device Risk Summary - WIDE TABLE
    elements.append(Paragraph("Device Risk Summary", heading_style))
    
    summary_data = [["Device", "Findings Count", "Risk Score"]]
    for _, row in summary_df.iterrows():
        # Use full device names - no truncation
        summary_data.append([str(row['Device']), str(row['Findings Count']), row['Risk Score']])
    
    # Calculate table width for landscape (11 inches wide - margins)
    table_width = landscape(letter)[0] - 0.6*inch  # 11 - 0.6 = 10.4 inches
    
    # Make Device column much wider to fit full names
    col_widths = [
        table_width * 0.70,  # Device - 70% of width for full names
        table_width * 0.15,  # Findings Count - 15%
        table_width * 0.15   # Risk Score - 15%
    ]
    
    summary_table = Table(summary_data, colWidths=col_widths, repeatRows=1)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4CAF50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),    # Device left-aligned
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'), # Counts and Risk centered
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),  # Enable word wrap for long device names
    ]))
    
    elements.append(summary_table)
    elements.append(Spacer(1, 20))

    # 2. Risk Distribution Chart
    elements.append(Paragraph("Risk Distribution (Devices by Risk Level)", heading_style))
    
    order = ["No Risk", "Low", "Medium", "High"]
    counts = [risk_counts.get(x, 0) for x in order]
    
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(order, counts, color=["lightgrey", "lightgreen", "gold", "crimson"])
    
    for bar, count in zip(bars, counts):
        if count > 0:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                   f'{count}', ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    ax.set_ylabel("Number of Devices", fontsize=12)
    ax.set_title("Device Risk Distribution", fontsize=14)
    
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        chart_path1 = tmp_file.name
    fig.savefig(chart_path1, bbox_inches='tight', dpi=120)
    plt.close(fig)
    
    elements.append(Image(chart_path1, width=7*inch, height=3.5*inch))
    elements.append(Spacer(1, 20))

    # 3. Findings by Category Chart
    elements.append(Paragraph("Findings by Category", heading_style))
    
    cat_names = list(category_counts.keys())
    cat_vals = list(category_counts.values())
    
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(cat_names, cat_vals, color="steelblue")
    
    for bar, count in zip(bars, cat_vals):
        if count > 0:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                   f'{count}', ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    ax.set_ylabel("Number of Findings", fontsize=12)
    ax.set_title("Findings Distribution per Category", fontsize=14)
    plt.xticks(rotation=45)
    
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        chart_path2 = tmp_file.name
    fig.savefig(chart_path2, bbox_inches='tight', dpi=120)
    plt.close(fig)
    
    elements.append(Image(chart_path2, width=8*inch, height=4*inch))
    elements.append(PageBreak())

    # 4. Detailed Findings with proper text wrapping
    elements.append(Paragraph("Detailed Findings", heading_style))
    
    if not df_findings.empty:
        devices = df_findings['File'].unique()
        
        for i, device in enumerate(devices):
            device_findings = df_findings[df_findings['File'] == device]
            
            elements.append(Paragraph(f"Device: {device}", device_style))
            elements.append(Spacer(1, 8))
            
            table_data = []
            
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=table_style,
                fontSize=7,
                fontName='Helvetica-Bold',
                textColor=colors.white,
                alignment=1
            )
            
            header_cells = [
                Paragraph("Category", header_style),
                Paragraph("Finding", header_style),
                Paragraph("Risk Description", header_style),
                Paragraph("Recommendation", header_style)
            ]
            table_data.append(header_cells)
            
            for _, finding in device_findings.iterrows():
                category_cell = Paragraph(str(finding['Category']), table_style)
                finding_cell = Paragraph(str(finding['Finding']), table_style)
                risk_cell = Paragraph(str(finding['RiskDesc']), table_style)
                recommendation_cell = Paragraph(str(finding['Recommendation']), table_style)
                
                table_data.append([category_cell, finding_cell, risk_cell, recommendation_cell])
            
            # Use full landscape width for detailed findings table
            table_width = landscape(letter)[0] - 0.6*inch
            col_widths = [
                table_width * 0.15,  # Category
                table_width * 0.20,  # Finding
                table_width * 0.30,  # Risk Description
                table_width * 0.35   # Recommendation
            ]
            
            device_table = Table(table_data, colWidths=col_widths, repeatRows=1)
            
            device_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2196F3')),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            
            elements.append(device_table)
            elements.append(Spacer(1, 15))
            
            if (i + 1) % 2 == 0 and (i + 1) < len(devices):
                elements.append(PageBreak())
    
    else:
        elements.append(Paragraph("No findings to report.", table_style))

    # Build PDF
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    # Cleanup
    try:
        os.unlink(chart_path1)
        os.unlink(chart_path2)
    except:
        pass
    
    return pdf_bytes

def generate_word_report(summary_df, df_findings, risk_counts, category_counts):
    # ... (include all your existing Word report generation code here)
    doc = Document()
    
    # Set wider margins for Word document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    
    # Title
    title = doc.add_heading('Network Configuration Audit Report', 0)
    doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")}')
    doc.add_paragraph()
    
    # 1. Device Risk Summary - WIDE TABLE
    doc.add_heading('Device Risk Summary', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False  # Disable autofit to control column widths
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Device'
    hdr_cells[1].text = 'Findings Count'
    hdr_cells[2].text = 'Risk Score'
    
    # Set column widths for Word table (wider first column)
    table.columns[0].width = Inches(6.0)  # Wide column for device names
    table.columns[1].width = Inches(1.5)  # Narrower for counts
    table.columns[2].width = Inches(1.5)  # Narrower for risk scores
    
    for _, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Device'])  # Full device name
        row_cells[1].text = str(row['Findings Count'])
        row_cells[2].text = str(row['Risk Score'])
    
    doc.add_paragraph()
    
    # 2. Risk Distribution
    doc.add_heading('Risk Distribution', level=1)
    for risk_level in ["No Risk", "Low", "Medium", "High"]:
        count = risk_counts.get(risk_level, 0)
        doc.add_paragraph(f'{risk_level}: {count} devices', style='List Bullet')
    
    doc.add_paragraph()
    
    # 3. Findings by Category
    doc.add_heading('Findings by Category', level=1)
    for category, count in category_counts.items():
        doc.add_paragraph(f'{category}: {count} findings', style='List Bullet')
    
    doc.add_paragraph()
    
    # 4. Detailed Findings
    doc.add_heading('Detailed Findings', level=1)
    
    if not df_findings.empty:
        devices = df_findings['File'].unique()
        
        for device in devices:
            doc.add_heading(f'Device: {device}', level=2)
            device_findings = df_findings[df_findings['File'] == device]
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Category'
            hdr_cells[1].text = 'Finding'
            hdr_cells[2].text = 'Risk Description'
            hdr_cells[3].text = 'Recommendation'
            
            # Set wider columns for Word
            table.columns[0].width = Inches(1.2)
            table.columns[1].width = Inches(2.0)
            table.columns[2].width = Inches(3.0)
            table.columns[3].width = Inches(3.0)
            
            for _, finding in device_findings.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(finding['Category'])
                row_cells[1].text = str(finding['Finding'])
                row_cells[2].text = str(finding['RiskDesc'])
                row_cells[3].text = str(finding['Recommendation'])
            
            doc.add_paragraph()
    else:
        doc.add_paragraph('No findings to report.')
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    word_bytes = buffer.getvalue()
    buffer.close()
    
    return word_bytes

def network_config_audit():
    st.title("🔐 Network Config Auditor")
    
    st.markdown("""
    ### Network Configuration Audit Input
    Upload individual configuration files in text format for security assessment and compliance auditing.

    **Current Input Requirements:**
    - Text files with .txt extension
    - Supports Individual and multiple file uploads

    **Audit Outputs:**
    - Detailed vulnerability findings with remediation guidance
    - Device-level risk scoring and categorization
    - Interactive security posture dashboard
    - Risk distribution heatmaps across security categories
    - Comprehensive export formats for reporting (CSV/PDF/DOCX)
    """)

    uploaded_files = st.file_uploader(
        "Select configuration files (.txt format only) — multiple selection enabled", 
        accept_multiple_files=True, 
        type=["txt"]
    )

    if uploaded_files:
        results = []  # list of tuples: (Finding, File, RiskDesc, Recommendation, Category)
        device_summary = defaultdict(list)

        def process_file_bytes(fname, raw_bytes):
            try:
                content = raw_bytes.decode("utf-8", errors="ignore")
            except Exception:
                content = raw_bytes.decode("latin-1", errors="ignore")
            file_findings = audit_config(fname, content)
            for f in file_findings:
                # f is (Finding, filename, RiskDesc, Recommendation, Category)
                results.append(f)
                device_summary[f[1]].append(f)
            return

        for uploaded in uploaded_files:
            name = uploaded.name
            lower = name.lower()
            # ZIP
            if lower.endswith(".zip"):
                try:
                    with zipfile.ZipFile(io.BytesIO(uploaded.read())) as zf:
                        for inner in zf.namelist():
                            if inner.endswith("/"):
                                continue
                            with zf.open(inner) as f:
                                raw = f.read()
                                process_file_bytes(inner, raw)
                except Exception as e:
                    st.warning(f"Failed to process ZIP {name}: {e}")

            # RAR
            elif lower.endswith(".rar"):
                try:
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".rar")
                    tmp.write(uploaded.read())
                    tmp.close()
                    with rarfile.RarFile(tmp.name) as rf:
                        for inner in rf.namelist():
                            if inner.endswith("/"):
                                continue
                            with rf.open(inner) as f:
                                raw = f.read()
                                process_file_bytes(inner, raw)
                    try:
                        os.remove(tmp.name)
                    except Exception:
                        pass
                except Exception as e:
                    st.warning(f"Failed to process RAR {name}: {e}")

            # Plain file (including no-extension)
            else:
                try:
                    raw = uploaded.read()
                    process_file_bytes(name, raw)
                except Exception as e:
                    st.warning(f"Failed to read file {name}: {e}")

        # show outputs
        if results:
            # build dataframe
            df = pd.DataFrame(results, columns=["Finding","File","RiskDesc","Recommendation","Category"])

            # Detailed findings view
            st.subheader("📋 Detailed Findings")
            st.dataframe(df[["File","Category","Finding","RiskDesc","Recommendation"]], width='stretch', height=320)

            # Device summary with risk score
            summary_rows = []
            for device, items in device_summary.items():
                score = get_risk_score(len(items))
                summary_rows.append((device, len(items), score))
            summary_df = pd.DataFrame(summary_rows, columns=["Device","Findings Count","Risk Score"])
            st.subheader("📊 Device Risk Summary (color-coded)")

            def color_row(r):
                score = r["Risk Score"]
                if score == "High":
                    return ['background-color:crimson;color:white']*3
                if score == "Medium":
                    return ['background-color:gold;color:black']*3
                if score == "Low":
                    return ['background-color:lightgreen;color:black']*3
                return ['background-color:lightgrey;color:black']*3

            st.dataframe(summary_df.style.apply(lambda row: color_row(row), axis=1), width='stretch', height=220)

            # Risk distribution chart
            st.subheader("📈 Risk Distribution")
            rc = summary_df["Risk Score"].value_counts().to_dict()
            order = ["No Risk","Low","Medium","High"]
            rc_plot = [rc.get(k,0) for k in order]
            fig, ax = plt.subplots()
            bars = ax.bar(order, rc_plot, color=["lightgrey","lightgreen","gold","crimson"])
            
            for bar, count in zip(bars, rc_plot):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{count}', ha='center', va='bottom', fontweight='bold')
            
            ax.set_ylabel("Number of Devices")
            ax.set_title("Device Risk Distribution")
            st.pyplot(fig)

            # Findings by category chart for Streamlit
            st.subheader("📊 Findings by Category")
            category_counts = df['Category'].value_counts().to_dict()
            cat_names = list(category_counts.keys())
            cat_vals = list(category_counts.values())
            fig2, ax2 = plt.subplots()
            bars2 = ax2.bar(cat_names, cat_vals, color="steelblue")
            
            for bar, count in zip(bars2, cat_vals):
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{count}', ha='center', va='bottom', fontweight='bold')
            
            ax2.set_ylabel("Number of Findings")
            ax2.set_title("Findings Distribution per Category")
            plt.xticks(rotation=45, ha="right")
            st.pyplot(fig2)

            # Heatmap
            st.subheader("🔥 Risk Heatmap per Category")
            heatmap_fig = generate_heatmap_figure(df)
            st.pyplot(heatmap_fig)

            # Downloads: CSVs
            csv_bytes = df.to_csv(index=False).encode("utf-8")
            st.download_button("📥 Download Detailed Findings (CSV)", csv_bytes, file_name="network_detailed_findings.csv", mime="text/csv")

            csv_summary = summary_df.to_csv(index=False).encode("utf-8")
            st.download_button("📥 Download Device Summary (CSV)", csv_summary, file_name="network_device_summary.csv", mime="text/csv")

            # Management Report Generation (PDF or Word)
            st.subheader("📄 Management Report")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("Generate PDF Report"):
                    with st.spinner("Building PDF Report..."):
                        category_counts = df['Category'].value_counts().to_dict()
                        risk_counts = summary_df["Risk Score"].value_counts().to_dict()
                        pdf_bytes = generate_pdf_report(summary_df, df, risk_counts, category_counts)
                        st.success("PDF report generated successfully!")
                        st.download_button("📥 Download PDF Report", 
                                         data=pdf_bytes, 
                                         file_name="network_audit_report.pdf", 
                                         mime="application/pdf")
            
            with col2:
                if st.button("Generate Word Report"):
                    with st.spinner("Building Word Report..."):
                        category_counts = df['Category'].value_counts().to_dict()
                        risk_counts = summary_df["Risk Score"].value_counts().to_dict()
                        word_bytes = generate_word_report(summary_df, df, risk_counts, category_counts)
                        st.success("Word report generated successfully!")
                        st.download_button("📥 Download Word Report", 
                                         data=word_bytes, 
                                         file_name="network_audit_report.docx", 
                                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        else:
            st.success("✅ No findings identified in uploaded files.")
    else:
        st.info("Upload individual text files (.txt format) for configuration analysis.")

# =============================================================================
# IAM FUNCTIONS
# =============================================================================

@st.cache_data
def load_excel(file):
    return pd.read_excel(file)

def find_matching_rows(df, column_name, disengaged_staff_list, threshold=70):
    """Find matching rows in the uploaded file using fuzzy matching."""
    if column_name not in df.columns:
        st.error(f"Column '{column_name}' not found.")
        return pd.DataFrame()
    
    matched_rows = pd.concat([
        df[df[column_name].isin(
            [match for match, score in process.extract(name, df[column_name].tolist(), scorer=fuzz.token_sort_ratio) if score >= threshold]
        )]
        for name in disengaged_staff_list
    ]).drop_duplicates()
    
    return matched_rows

def iam_main_page():
    st.title("🏠 Identity Access Management Tool")
    
    # Initialize session state for matched results if not exists
    if "matched_results" not in st.session_state:
        st.session_state["matched_results"] = {}
    
    # Step 1: Upload Disengaged Staff List
    st.header("Step 1: Upload Disengaged Staff List")
    disengaged_file = st.file_uploader("📂 Upload an Excel file", type=["xlsx"], key="disengaged")
    disengaged_list = []
    
    if disengaged_file:
        disengaged_df = load_excel(disengaged_file)
        disengaged_column = st.selectbox("🛑 Select column with disengaged staff names", disengaged_df.columns)
        disengaged_list = disengaged_df[disengaged_column].dropna().tolist()
        st.success("✅ Disengaged staff list uploaded.")
    
    # Step 2: Upload System Users List
    st.header("Step 2: Upload System Users List")
    app_file = st.file_uploader("📂 Upload an Excel file", type=["xlsx"], key="app")
    
    if app_file:
        app_df = load_excel(app_file)
        app_column = st.selectbox("🔎 Select column to match", app_df.columns, key="app_col")
        app_name = st.text_input("🖥️ Enter the system name", key="app_name")
        
        if st.button("🔍 Run Matching"):
            if app_name and disengaged_list:
                matched_df = find_matching_rows(app_df, app_column, disengaged_list)
                if not matched_df.empty:
                    st.session_state["matched_results"][app_name] = matched_df
                    st.success(f"✅ Matching completed for {app_name}.")
                else:
                    st.warning(f"No matches found for {app_name}.")
            else:
                st.warning("Please provide a system name and ensure the disengaged staff list is uploaded.")
            
            # Clear only the Step 2 fields by deleting their keys
            for key in ["app", "app_col", "app_name"]:
                if key in st.session_state:
                    del st.session_state[key]
            # No full rerun so that previous results remain intact.
    
    # Step 3: Download Consolidated Results
    # Show this step if there are any matched results.
    if st.session_state["matched_results"]:
        st.header("Step 3: Review and Download Results")
        
        # Show summary of matched results
        summary_df = pd.DataFrame([
            {"System": app, "Matches": len(data)} 
            for app, data in st.session_state["matched_results"].items()
        ])
        st.dataframe(summary_df)

        selected_app = st.selectbox("🔍 Select system to preview", list(st.session_state["matched_results"].keys()))
        if selected_app:
            st.dataframe(st.session_state["matched_results"][selected_app])

        # Consolidate results into one Excel file
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            for app, data in st.session_state["matched_results"].items():
                data.to_excel(writer, sheet_name=app, index=False)
        output.seek(0)

        st.download_button(
            label="📥 Download Consolidated Results",
            data=output,
            file_name="Consolidated_Results.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

def duplicate_user_provisioning():
    st.title("🔁 Duplicate User Provisioning")
    uploaded_file = st.file_uploader("Upload System Users", type=["xls", "xlsx"])
    
    if not uploaded_file:
        st.info("Please upload an Excel file to proceed.")
        st.stop()
        
    sys_users = pd.read_excel(uploaded_file)
    username_column = st.selectbox("Select the column containing usernames", sys_users.columns)
    
    if not username_column:
        st.stop()

    # 1) Identify users with >1 provisioning
    vc = sys_users[username_column].value_counts()
    dup_users = vc[vc > 1].index.tolist()
    dup_df = sys_users[sys_users[username_column].isin(dup_users)]
    
    # 2) Show raw rows
    st.subheader("🔍 Raw Rows for Users with Multiple Provisions")
    st.dataframe(dup_df)
    
    # 3) Show summary
    summary = (
        dup_df[username_column]
        .value_counts()
        .reset_index()
        .rename(columns={"index": "Username", username_column: "Occurrences"})
    )
    st.subheader("📊 Occurrence Summary")
    st.dataframe(summary)
    
    # 4) Download the raw duplicate rows
    if not dup_df.empty:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
            dup_df.to_excel(writer, sheet_name="Multiple_Provisions", index=False)
        buffer.seek(0)
        
        st.download_button(
            label="📥 Download Users with Multiple Provisions",
            data=buffer,
            file_name="users_multiple_provisions.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.info("No users with multiple provisions to download.")

def database_groups():
    st.title("📂 Database Groups Management")
    uploaded_file = st.file_uploader("📂 Upload ORACLE DBA_USER REPORT", type=["xls", "xlsx"])
    st.markdown("""
    ### 📤 What to Upload
    Upload the **ORACLE DBA_USER report** exported from your database environment (xlsx format).  
    Ensure the file contains key user account details such as **Username**, **Account Status**, **Created Date**, **Profile**, **Password Versions**, **Privileges**.

    ### ⚙️ What the Tool Does
    Once uploaded, the tool automatically analyzes the report to:
    - Identify **inactive or unauthorized accounts**
    - Detect **default or weak credentials**
    - Highlight **users with excessive or unnecessary privileges**
    - Check for **Segregation of Duties (SoD)** violations
    - Assess compliance with **security and access control standards**

    💡 The tool provides a comprehensive security overview to help auditors validate **user access governance** and ensure **database integrity**.
    """)
    
    if uploaded_file:
        try:
            db_users = pd.read_excel(uploaded_file)
            
            # Display dataset info
            st.success(f"✅ Successfully loaded {len(db_users)} user accounts")
            st.info(f"📊 Columns detected: {', '.join(db_users.columns)}")
            
            # =============================================================================
            # COLUMN SELECTION SECTION
            # =============================================================================
            st.header("🔧 Configure Security Checks")
            st.markdown("Select the appropriate columns from your file for each security check:")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                username_col = st.selectbox(
                    "👤 Username Column",
                    options=[''] + list(db_users.columns),
                    index=0,
                    help="Select the column containing usernames"
                )
                
                status_col = st.selectbox(
                    "🔒 Account Status Column",
                    options=[''] + list(db_users.columns),
                    index=0,
                    help="Select the column showing account status (OPEN, LOCKED, EXPIRED, etc.)"
                )
            
            with col2:
                profile_col = st.selectbox(
                    "👥 Profile Column",
                    options=[''] + list(db_users.columns),
                    index=0,
                    help="Select the column containing user profiles"
                )
                
                created_col = st.selectbox(
                    "📅 Account Created Date Column",
                    options=[''] + list(db_users.columns),
                    index=0,
                    help="Select the column showing when accounts were created"
                )
            
            with col3:
                password_col = st.selectbox(
                    "🔐 Password Version Column",
                    options=[''] + list(db_users.columns),
                    index=0,
                    help="Select the column showing password versions (optional)"
                )
                
                privilege_col = st.selectbox(
                    "👑 Privilege/Role Column",
                    options=[''] + list(db_users.columns),
                    index=0,
                    help="Select the column showing user privileges or roles (optional)"
                )
            
            # =============================================================================
            # SECURITY ANALYSIS SECTION
            # =============================================================================
            st.header("🔍 Security Analysis Results")
            
            # Initialize findings
            security_findings = []
            analysis_results = {}
            
            # 1. Identify Inactive/Locked Accounts (INFORMATIONAL - Not a risk)
            if status_col:
                st.subheader("🔒 Account Status Analysis")
                try:
                    status_counts = db_users[status_col].value_counts()
                    st.dataframe(status_counts)
                    
                    # Identify problematic statuses
                    inactive_statuses = ['LOCKED', 'EXPIRED', 'EXPIRED(GRACE)', 'INACTIVE', 'LOCKED(TIMED)']
                    inactive_accounts = db_users[
                        db_users[status_col].astype(str).str.upper().isin([s.upper() for s in inactive_statuses])
                    ]
                    
                    if not inactive_accounts.empty:
                        st.info(f"ℹ️ {len(inactive_accounts)} inactive/locked accounts found (this is normal security practice)")
                        display_cols = [username_col, status_col]
                        if created_col:
                            display_cols.append(created_col)
                        if profile_col:
                            display_cols.append(profile_col)
                        
                        st.dataframe(inactive_accounts[display_cols].head(15))
                        analysis_results['inactive_accounts'] = inactive_accounts
                    else:
                        st.success("✅ No inactive/locked accounts found")
                        
                except Exception as e:
                    st.error(f"Error analyzing account status: {str(e)}")
            
            # 2. Detect Default/Weak Credentials
            if username_col:
                st.subheader("⚠️ Default Account Detection")
                default_users = ['SYS', 'SYSTEM', 'DBSNMP', 'OUTLN', 'MGMT_VIEW', 'SYSMAN', 'SCOTT']
                try:
                    found_default = db_users[
                        db_users[username_col].astype(str).str.upper().isin([u.upper() for u in default_users])
                    ]
                    
                    if not found_default.empty:
                        st.warning(f"⚠️ {len(found_default)} default database accounts found")
                        display_cols = [username_col]
                        if status_col:
                            display_cols.append(status_col)
                        if profile_col:
                            display_cols.append(profile_col)
                            
                        st.dataframe(found_default[display_cols])
                        security_findings.append(f"⚠️ {len(found_default)} default database accounts found")
                        analysis_results['default_accounts'] = found_default
                    else:
                        st.success("✅ No default database accounts found")
                        
                except Exception as e:
                    st.error(f"Error analyzing default accounts: {str(e)}")
            
            # 3. Identify Users with DBA Privileges
            if profile_col:
                st.subheader("👑 Privileged Account Analysis")
                dba_profiles = ['DBA', 'SYSDBA', 'SYSOPER', 'SYSDG', 'SYSBACKUP', 'SYSKM', 'SYSRAC', 'SYSASM']
                try:
                    dba_users = db_users[
                        db_users[profile_col].astype(str).str.upper().isin([p.upper() for p in dba_profiles])
                    ]
                    
                    if not dba_users.empty:
                        st.warning(f"👑 {len(dba_users)} users with DBA/privileged profiles found")
                        display_cols = [username_col, profile_col]
                        if status_col:
                            display_cols.append(status_col)
                            
                        st.dataframe(dba_users[display_cols])
                        security_findings.append(f"👑 {len(dba_users)} users with DBA/privileged profiles")
                        analysis_results['dba_users'] = dba_users
                    else:
                        st.success("✅ No users with DBA profiles found")
                        
                except Exception as e:
                    st.error(f"Error analyzing privileged accounts: {str(e)}")
            
            # 4. Profile Distribution Analysis
            if profile_col:
                st.subheader("📊 Profile Distribution")
                try:
                    profile_counts = db_users[profile_col].value_counts()
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.dataframe(profile_counts.head(15))
                    
                    with col2:
                        if len(profile_counts) > 0:
                            fig, ax = plt.subplots(figsize=(10, 6))
                            profile_counts.head(10).plot(kind='bar', ax=ax, color='skyblue')
                            ax.set_title('Top 10 Profiles by User Count')
                            ax.set_ylabel('Number of Users')
                            ax.set_xlabel('Profile Name')
                            plt.xticks(rotation=45, ha='right')
                            st.pyplot(fig)
                            
                    # Identify profiles with many users (potential risk)
                    large_profiles = profile_counts[profile_counts > 10]
                    if not large_profiles.empty:
                        security_findings.append(f"📊 {len(large_profiles)} profiles with more than 10 users")
                        
                except Exception as e:
                    st.error(f"Error analyzing profile distribution: {str(e)}")
            
            # 5. Account Age & Security Risk Analysis
            if created_col and status_col:
                st.subheader("📅 Account Age & Security Risk Analysis")
                try:
                    # Convert to datetime
                    db_users['CREATED_DATE'] = pd.to_datetime(db_users[created_col], errors='coerce')
                    valid_dates = db_users['CREATED_DATE'].notna()
                    
                    if valid_dates.any():
                        one_year_ago = pd.Timestamp.now() - pd.DateOffset(years=1)
                        
                        # Find accounts older than 1 year
                        old_accounts = db_users[valid_dates & (db_users['CREATED_DATE'] < one_year_ago)]
                        
                        if not old_accounts.empty:
                            # Define active statuses (accounts that are still usable)
                            active_statuses = ['OPEN', 'ACTIVE', 'VALID', 'ENABLED']
                            
                            # Find OLD accounts that are still ACTIVE - HIGH SECURITY RISK!
                            old_active_accounts = old_accounts[
                                old_accounts[status_col].astype(str).str.upper().isin([s.upper() for s in active_statuses])
                            ]
                            
                            # Find old accounts that are properly locked/expired
                            old_inactive_accounts = old_accounts[
                                ~old_accounts[status_col].astype(str).str.upper().isin([s.upper() for s in active_statuses])
                            ]
                            
                            # HIGH RISK: Old accounts still active
                            if not old_active_accounts.empty:
                                st.error(f"🚨 HIGH RISK: {len(old_active_accounts)} accounts older than 1 year are still ACTIVE!")
                                
                                display_cols = [username_col, status_col, 'CREATED_DATE']
                                if profile_col:
                                    display_cols.append(profile_col)
                                
                                st.dataframe(old_active_accounts[display_cols].sort_values('CREATED_DATE'))
                                
                                # Show risk breakdown by profile
                                if profile_col:
                                    risk_by_profile = old_active_accounts[profile_col].value_counts()
                                    if not risk_by_profile.empty:
                                        st.write("**High-risk accounts by profile:**")
                                        st.dataframe(risk_by_profile)
                                
                                security_findings.append(f"🚨 HIGH RISK: {len(old_active_accounts)} old accounts (>1 year) still active")
                                analysis_results['old_active_accounts_high_risk'] = old_active_accounts
                            
                            # Informational: Old accounts that are properly managed
                            if not old_inactive_accounts.empty:
                                st.success(f"✅ {len(old_inactive_accounts)} old accounts are properly locked/expired")
                                analysis_results['old_inactive_accounts'] = old_inactive_accounts
                            
                            # Show overall old account statistics
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Total Old Accounts", len(old_accounts))
                            with col2:
                                st.metric("High Risk (Active)", len(old_active_accounts), delta=f"{len(old_active_accounts)/len(old_accounts)*100:.1f}%", delta_color="inverse")
                            with col3:
                                st.metric("Properly Managed", len(old_inactive_accounts), delta=f"{len(old_inactive_accounts)/len(old_accounts)*100:.1f}%")
                            
                            # Show creation timeline of high-risk accounts
                            if not old_active_accounts.empty:
                                st.subheader("📈 High-Risk Account Creation Timeline")
                                try:
                                    # Group by year-month
                                    old_active_accounts['CREATED_YEAR_MONTH'] = old_active_accounts['CREATED_DATE'].dt.to_period('M')
                                    timeline_data = old_active_accounts['CREATED_YEAR_MONTH'].value_counts().sort_index()
                                    
                                    fig, ax = plt.subplots(figsize=(12, 6))
                                    timeline_data.plot(kind='bar', ax=ax, color='red', alpha=0.7)
                                    ax.set_title('High-Risk Active Accounts by Creation Date\n(Accounts >1 year old still active)')
                                    ax.set_ylabel('Number of High-Risk Accounts')
                                    ax.set_xlabel('Creation Period')
                                    plt.xticks(rotation=45, ha='right')
                                    st.pyplot(fig)
                                except Exception as timeline_error:
                                    st.warning("Could not generate timeline chart")
                        
                        else:
                            st.success("✅ No accounts older than 1 year found")
                            
                        # Show recent account statistics for comparison
                        recent_accounts = db_users[valid_dates & (db_users['CREATED_DATE'] >= one_year_ago)]
                        st.success(f"🆕 {len(recent_accounts)} accounts created in the last year")
                        
                    else:
                        st.warning("Could not parse creation dates from selected column")
                        
                except Exception as e:
                    st.error(f"Error analyzing account ages: {str(e)}")
            
            # 6. Password Policy Analysis
            if password_col:
                st.subheader("🔐 Password Security Analysis")
                try:
                    pwd_versions = db_users[password_col].value_counts()
                    st.dataframe(pwd_versions)
                    
                    # Check for outdated password versions
                    outdated_pwd = ['10G', '11G']  # Add versions you consider outdated
                    users_outdated_pwd = db_users[
                        db_users[password_col].astype(str).str.upper().isin([v.upper() for v in outdated_pwd])
                    ]
                    
                    if not users_outdated_pwd.empty:
                        st.warning(f"🔐 {len(users_outdated_pwd)} users using older password versions")
                        security_findings.append(f"🔐 {len(users_outdated_pwd)} users using older password versions")
                        analysis_results['outdated_password_users'] = users_outdated_pwd
                        
                except Exception as e:
                    st.error(f"Error analyzing password versions: {str(e)}")
            
            # 7. Database Group Integrity Analysis
            st.subheader("🏗️ Database Group & Default User Integrity")

            # Check if database has groups/profiles
            if profile_col:
                try:
                    total_profiles = db_users[profile_col].nunique()
                    total_users = len(db_users)
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.metric("Total Profiles/Groups", total_profiles)
                    with col2:
                        st.metric("Total Users", total_users)
                    
                    # Check 1: Database has proper group structure
                    if total_profiles <= 1:
                        st.error("🚨 CRITICAL: Database has only 1 profile/group - No proper segregation of duties!")
                        security_findings.append("🚨 CRITICAL: Only 1 profile/group found - No SoD implementation")
                    elif total_profiles < 5:
                        st.warning("⚠️ WARNING: Limited profile/group structure - Consider implementing more granular access controls")
                        security_findings.append("⚠️ Limited profile/group structure - Consider more granular controls")
                    else:
                        st.success(f"✅ Good: Database has {total_profiles} profiles/groups for proper access control")
                    
                    # Check 2: Default profiles analysis
                    default_profiles = ['DEFAULT', 'BASIC', 'STANDARD', 'NONE']
                    default_profile_users = db_users[
                        db_users[profile_col].astype(str).str.upper().isin([p.upper() for p in default_profiles])
                    ]
                    
                    if not default_profile_users.empty:
                        st.warning(f"⚠️ {len(default_profile_users)} users assigned to default profiles")
                        
                        # Show default profile users
                        display_cols = [username_col, profile_col]
                        if status_col:
                            display_cols.append(status_col)
                        
                        st.dataframe(default_profile_users[display_cols])
                        
                        # Check if non-service accounts are in default profiles
                        service_accounts = ['SYS', 'SYSTEM', 'DBSNMP', 'ORACLE_OCM', 'XS$NULL']
                        non_service_in_default = default_profile_users[
                            ~default_profile_users[username_col].astype(str).str.upper().isin([s.upper() for s in service_accounts])
                        ]
                        
                        if not non_service_in_default.empty:
                            st.error(f"🚨 HIGH RISK: {len(non_service_in_default)} NON-SERVICE accounts in default profiles!")
                            st.dataframe(non_service_in_default[display_cols])
                            security_findings.append(f"🚨 HIGH RISK: {len(non_service_in_default)} non-service accounts in default profiles")
                        else:
                            st.info("✅ Only service accounts in default profiles (acceptable)")
                    
                except Exception as e:
                    st.error(f"Error analyzing group structure: {str(e)}")

            # 8. Default User Privilege & Integrity Analysis
            if username_col:
                st.subheader("👑 Default User Security Analysis")
                
                try:
                    # Define default Oracle accounts and their expected states
                    default_account_checks = {
                        'SYS': {
                            'expected_status': 'OPEN',
                            'risk_if': 'LOCKED',  # SYS should generally be open
                            'description': 'Data dictionary owner - critical system account'
                        },
                        'SYSTEM': {
                            'expected_status': 'OPEN', 
                            'risk_if': 'LOCKED',
                            'description': 'Administrative operations - should be open'
                        },
                        'DBSNMP': {
                            'expected_status': 'OPEN',
                            'risk_if': 'LOCKED', 
                            'description': 'Enterprise Manager agent - should be open'
                        },
                        'OUTLN': {
                            'expected_status': 'OPEN',
                            'risk_if': 'LOCKED',
                            'description': 'Plan stability - can be locked if not used'
                        },
                        'MGMT_VIEW': {
                            'expected_status': 'OPEN',
                            'risk_if': 'LOCKED',
                            'description': 'Enterprise Manager - should be open if EM used'
                        },
                        'SYSMAN': {
                            'expected_status': 'OPEN', 
                            'risk_if': 'LOCKED',
                            'description': 'Enterprise Manager super admin - should be open if EM used'
                        },
                        'SCOTT': {
                            'expected_status': 'LOCKED',
                            'risk_if': 'OPEN',
                            'description': 'Sample/training account - should be LOCKED in production'
                        },
                        'HR': {
                            'expected_status': 'LOCKED', 
                            'risk_if': 'OPEN',
                            'description': 'Sample/training account - should be LOCKED in production'
                        },
                        'OE': {
                            'expected_status': 'LOCKED',
                            'risk_if': 'OPEN', 
                            'description': 'Sample/training account - should be LOCKED in production'
                        }
                    }
                    
                    default_user_issues = []
                    default_user_summary = []
                    
                    for default_user, expected_config in default_account_checks.items():
                        user_data = db_users[db_users[username_col].astype(str).str.upper() == default_user.upper()]
                        
                        if not user_data.empty:
                            actual_status = user_data[status_col].iloc[0] if status_col else 'UNKNOWN'
                            profile_name = user_data[profile_col].iloc[0] if profile_col else 'UNKNOWN'
                            
                            status_check = "✅" if str(actual_status).upper() == expected_config['expected_status'].upper() else "❌"
                            
                            # Check if account is in risky state
                            is_risky = str(actual_status).upper() == expected_config['risk_if'].upper()
                            
                            default_user_summary.append({
                                'Username': default_user,
                                'Found': 'YES',
                                'Current Status': actual_status,
                                'Expected Status': expected_config['expected_status'],
                                'Status Check': status_check,
                                'Risk': 'HIGH' if is_risky else 'LOW',
                                'Description': expected_config['description']
                            })
                            
                            if is_risky:
                                default_user_issues.append(f"🚨 {default_user} is {actual_status} but should be {expected_config['expected_status']} - {expected_config['description']}")
                        
                        else:
                            default_user_summary.append({
                                'Username': default_user,
                                'Found': 'NO',
                                'Current Status': 'NOT FOUND',
                                'Expected Status': 'N/A',
                                'Status Check': '⚠️',
                                'Risk': 'MEDIUM',
                                'Description': expected_config['description']
                            })
                            default_user_issues.append(f"⚠️ Default account {default_user} not found in database")
                    
                    # Display default user analysis
                    if default_user_summary:
                        summary_df = pd.DataFrame(default_user_summary)
                        st.dataframe(summary_df)
                        
                        # Show risk summary
                        high_risk_count = len([u for u in default_user_issues if '🚨' in u])
                        medium_risk_count = len([u for u in default_user_issues if '⚠️' in u])
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("High Risk Default Users", high_risk_count)
                        with col2:
                            st.metric("Medium Risk Findings", medium_risk_count)
                        
                        # Add findings to security report
                        security_findings.extend(default_user_issues)
                        analysis_results['default_user_analysis'] = pd.DataFrame(default_user_summary)
                    
                except Exception as e:
                    st.error(f"Error analyzing default users: {str(e)}")

            # 9. Orphaned/Unauthorized Profile Analysis
            if profile_col and username_col:
                st.subheader("🔍 Profile Usage & Authorization Analysis")
                
                try:
                    # Check for profiles with very few users (potential orphaned/custom profiles)
                    profile_user_counts = db_users[profile_col].value_counts()
                    
                    # Profiles with only 1 user (potential service accounts or custom profiles)
                    single_user_profiles = profile_user_counts[profile_user_counts == 1]
                    
                    if not single_user_profiles.empty:
                        st.info(f"🔍 {len(single_user_profiles)} profiles have only 1 user")
                        
                        # Get the actual users for these profiles
                        single_user_details = db_users[db_users[profile_col].isin(single_user_profiles.index)]
                        
                        display_cols = [username_col, profile_col]
                        if status_col:
                            display_cols.append(status_col)
                        if created_col:
                            display_cols.append(created_col)
                        
                        st.dataframe(single_user_details[display_cols].head(10))
                        
                        # Check if these are known service accounts or potential unauthorized profiles
                        service_accounts = ['SYS', 'SYSTEM', 'DBSNMP', 'SYSMAN', 'ORACLE_OCM']
                        unknown_single_users = single_user_details[
                            ~single_user_details[username_col].astype(str).str.upper().isin([s.upper() for s in service_accounts])
                        ]
                        
                        if not unknown_single_users.empty:
                            st.warning(f"⚠️ {len(unknown_single_users)} non-service accounts in single-user profiles - review for authorization")
                            security_findings.append(f"⚠️ {len(unknown_single_users)} non-service accounts in single-user profiles - potential unauthorized access")
                    
                    # Check for profiles with no users (orphaned profiles)
                    # This would require comparing against all possible profiles in the database
                    
                except Exception as e:
                    st.error(f"Error analyzing profile usage: {str(e)}")

            # 10. Privilege Escalation Risk Analysis
            if profile_col and username_col:
                st.subheader("⚡ Privilege Escalation Risk Analysis")
                
                try:
                    # Identify users with powerful profiles but non-standard names
                    powerful_profiles = ['DBA', 'SYSDBA', 'SYSOPER', 'SYSDG', 'SYSBACKUP', 'SYSKM']
                    
                    powerful_users = db_users[
                        db_users[profile_col].astype(str).str.upper().isin([p.upper() for p in powerful_profiles])
                    ]
                    
                    if not powerful_users.empty:
                        # Check if these are known administrative accounts
                        known_admin_accounts = ['SYS', 'SYSTEM', 'SYSMAN']
                        
                        unknown_powerful_users = powerful_users[
                            ~powerful_users[username_col].astype(str).str.upper().isin([a.upper() for a in known_admin_accounts])
                        ]
                        
                        if not unknown_powerful_users.empty:
                            st.error(f"🚨 HIGH RISK: {len(unknown_powerful_users)} non-standard users with powerful privileges!")
                            
                            display_cols = [username_col, profile_col]
                            if status_col:
                                display_cols.append(status_col)
                            if created_col:
                                display_cols.append(created_col)
                            
                            st.dataframe(unknown_powerful_users[display_cols])
                            security_findings.append(f"🚨 HIGH RISK: {len(unknown_powerful_users)} non-standard users with powerful admin privileges")
                            
                            # Check if any of these are active
                            if status_col:
                                active_powerful_unknown = unknown_powerful_users[
                                    unknown_powerful_users[status_col].astype(str).str.upper().isin(['OPEN', 'ACTIVE'])
                                ]
                                if not active_powerful_unknown.empty:
                                    st.error(f"🔴 CRITICAL: {len(active_powerful_unknown)} unknown users with admin privileges are ACTIVE!")
                                    security_findings.append(f"🔴 CRITICAL: {len(active_powerful_unknown)} unknown admin users are ACTIVE")
                        else:
                            st.success("✅ All powerful privileges assigned to known administrative accounts")
                    
                except Exception as e:
                    st.error(f"Error analyzing privilege escalation risks: {str(e)}")

            # =============================================================================
            # ORIGINAL GROUP MANAGEMENT FUNCTIONALITY
            # =============================================================================
            if profile_col:
                st.header("👥 Profile Management")
                
                # Extract Unique Profiles 
                unique_profiles = db_users[profile_col].unique()

                # Select the profile to View its Users 
                selected_profile = st.selectbox("🔎 Select a Profile Name: ", unique_profiles)

                # Display Users for the Selected Profile
                profile_users = db_users[db_users[profile_col] == selected_profile]
                st.subheader(f"🗂 Users with Profile: **{selected_profile}**")
                st.dataframe(profile_users)

                # Consolidate all profile users into one Excel file with separate sheets 
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                    for profile in unique_profiles:
                        group = db_users[db_users[profile_col] == profile]
                        # Limit sheet name to 31 characters (Excel limitations)
                        sheet_name = str(profile)[:31]
                        group.to_excel(writer, sheet_name=sheet_name, index=False)
                output.seek(0)

                st.download_button(
                    label = "📥 Download Consolidated Users of Profiles", 
                    data = output, 
                    file_name="Consolidated_Users_of_Profiles.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
            # =============================================================================
            # SECURITY FINDINGS SUMMARY
            # =============================================================================
            st.header("📋 Security Findings Summary")
            
            if security_findings:
                for i, finding in enumerate(security_findings, 1):
                    st.write(f"{i}. {finding}")
                
                # Risk Level Assessment
                risk_count = len(security_findings)
                critical_count = len([f for f in security_findings if '🔴' in f])
                high_count = len([f for f in security_findings if '🚨' in f and '🔴' not in f])
                
                if critical_count > 0:
                    st.error("🔴 CRITICAL RISK: Immediate action required!")
                elif high_count > 2:
                    st.error("🚨 HIGH RISK: Multiple serious security issues detected")
                elif risk_count > 5:
                    st.warning("🟡 MEDIUM RISK: Several security issues detected")
                else:
                    st.info("🟢 LOW RISK: Minor security issues detected")
            else:
                st.success("🎉 EXCELLENT: No security issues detected!")
            
            # =============================================================================
            # EXPORT SECTION
            # =============================================================================
            st.header("📤 Export Results")
            
            if st.button("📊 Generate Comprehensive Audit Report"):
                with st.spinner("Generating audit report..."):
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                        # Sheet 1: All Users
                        db_users.to_excel(writer, sheet_name='All_Users', index=False)
                        
                        # Sheet 2: Security Findings Summary
                        findings_df = pd.DataFrame({
                            'Finding': security_findings,
                            'Risk_Level': ['Critical' if '🔴' in f else 'High' if '🚨' in f else 'Medium' if '⚠️' in f else 'Low' for f in security_findings],
                            'Timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        })
                        findings_df.to_excel(writer, sheet_name='Security_Findings', index=False)
                        
                        # Additional sheets for each analysis result
                        for result_name, result_data in analysis_results.items():
                            if not result_data.empty:
                                sheet_name = result_name.replace('_', ' ').title()[:31]
                                result_data.to_excel(writer, sheet_name=sheet_name, index=False)
                        
                        # Profile Summary
                        if profile_col:
                            profile_summary = db_users[profile_col].value_counts().reset_index()
                            profile_summary.columns = ['Profile', 'User_Count']
                            profile_summary.to_excel(writer, sheet_name='Profile_Summary', index=False)
                    
                    output.seek(0)
                    
                    st.download_button(
                        label="📥 Download Comprehensive Audit Report",
                        data=output,
                        file_name=f"Database_Security_Audit_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="audit_report"
                    )
            
            # Quick CSV export of findings
            if security_findings:
                csv_findings = pd.DataFrame({
                    'Security_Finding': security_findings,
                    'Timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                })
                csv_output = csv_findings.to_csv(index=False).encode('utf-8')
                
                st.download_button(
                    label="📋 Download Security Findings (CSV)",
                    data=csv_output,
                    file_name="security_findings.csv",
                    mime="text/csv",
                    key="findings_csv"
                )
                
        except Exception as e:
            st.error(f"❌ Error processing file: {str(e)}")
            st.info("Please ensure the file is a valid Excel file with proper data formatting.")
    else:
        # Show sample expected columns when no file is uploaded
        st.info("""
        **📋 Expected Columns in DBA_USER Report:**
        - **Username Column**: Contains user identifiers (e.g., USERNAME, USER_NAME, USER)
        - **Status Column**: Shows account status (e.g., ACCOUNT_STATUS, STATUS, USER_STATUS)
        - **Profile Column**: Contains profile assignments (e.g., PROFILE, USER_PROFILE)
        - **Created Date**: Account creation date (e.g., CREATED, CREATE_DATE, DATE_CREATED)
        - **Password Version**: Password hash versions (e.g., PASSWORD_VERSIONS, PWD_VERSION)
        - **Privilege/Role**: User roles and privileges (e.g., PRIVILEGE, ROLE, GRANTED_ROLE)
        
        **💡 Tip**: The tool will automatically detect your column names and let you map them to the required fields.
        """)

def database_privilege_users():
    st.title("🔑 Database Privilege Users")
    uploaded_file = st.file_uploader("📂 Upload ORACLE DBA_ROLE_PRIVS", type=["xlsx"], key="db_priv")
    st.markdown("""
                """)
    if uploaded_file:
        db_priv_df = pd.read_excel(uploaded_file)

        # Extract Unique Admin Options 
        unique_admin_names = db_priv_df['ADMIN OPTION'].unique()
        
        # Select an Admin Option to View its Users
        selected_admin = st.selectbox("🔎 Select Admin Option", db_priv_df["ADMIN OPTION"].unique())

        # Display Users for the Selected Resource 
        admin = db_priv_df.groupby('ADMIN OPTION')
        users_df = admin.get_group(selected_admin)
        st.subheader(f"Users for Admin Option: **{selected_admin}**")
        st.dataframe(users_df)

        # 📤 Consolidate all resource groups into one Excel file with separate sheets
        import io
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            for users in unique_admin_names:
                group = db_priv_df[db_priv_df["ADMIN OPTION"] == users]
                # Limit sheet name to 31 characters (Excel limitation)
                sheet_name = users[:31]
                group.to_excel(writer, sheet_name=sheet_name, index=False)
        output.seek(0)

        st.download_button(
            label="📥 Download Users in Admin Options",
            data=output,
            file_name="Consolidated_Admin_Options_Users.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

def database_profiles():
    st.title("🗂 Database Profiles")
    uploaded_file = st.file_uploader("📂 Upload ORACLE_DBA_PROFILES", type=["xls", "xlsx"])

    if uploaded_file:
        database_profile = pd.read_excel(uploaded_file)

        # 🎯 Extract Unique Resource Names
        unique_resource_names = database_profile['RESOURCE NAME'].unique()

        # 🔍 Select a Resource Name to View its Users
        selected_resource = st.selectbox("🔎 Select a Resource Name:", unique_resource_names)

        # 🎲 Display Users for the Selected Resource
        profiles = database_profile.groupby('RESOURCE NAME')
        users_df = profiles.get_group(selected_resource)
        st.subheader(f"🗂 Profiles for: **{selected_resource}**")
        st.dataframe(users_df)

        # 📤 Consolidate all resource groups into one Excel file with separate sheets
        import io
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            for resource in unique_resource_names:
                group = database_profile[database_profile["RESOURCE NAME"] == resource]
                # Limit sheet name to 31 characters (Excel limitation)
                sheet_name = resource[:31]
                group.to_excel(writer, sheet_name=sheet_name, index=False)
        output.seek(0)

        st.download_button(
            label="📥 Download Consolidated Profiles",
            data=output,
            file_name="Consolidated_Profiles.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# =============================================================================
# MAIN APPLICATION ROUTING
# =============================================================================

def main():
    # Main title
    st.title("🔐 Your-IT-Auditor")
    st.markdown("---")
    
    # Route to appropriate page based on selection
    if main_category == "🔐 Network Security":
        if page == "📅 Audit Planner":
            audit_planner()
        elif page == "📊 Config Audit":
            network_config_audit()
    else:  # Identity & Access Management
        if page == "🏠 IAM Main":
            iam_main_page()
        elif page == "🔁 Duplicate User Provisioning":
            duplicate_user_provisioning()
        elif page == "📂 Database Groups":
            database_groups()
        elif page == "🔑 Database Privilege Users":
            database_privilege_users()
        elif page == "🗂 Database Profiles":
            database_profiles()

if __name__ == "__main__":
    main()